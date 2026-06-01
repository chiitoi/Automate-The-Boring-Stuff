#Chapter 17 of Automate the Boring Stuff by Al Sweigart
#https://automatetheboringstuff.com/3e/chapter17.html

import docx

def bold_words(filename, word):
    doc = docx.Document(filename)
    new_doc = docx.Document()

    #iterates over each paragraph in the original document
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        new_para = new_doc.add_paragraph()

        #we modify the original paragraph in each iteration
        while(len(paragraph_text) > 0):
            if word in paragraph_text:
                index1 = paragraph_text.find(word)
                
                #splice the string until we find the index where the word is
                #create a separate run with separate font style
                if index1 == 0:
                    #print(paragraph_text[0:len(word)])
                    new_run = new_para.add_run(paragraph_text[0:len(word)])
                    new_run.font.bold = True
                    paragraph_text = paragraph_text[len(word):]

                #splice the string so that we add everything before the word as a separate run
                else:
                    #print(paragraph_text[0:index1])
                    new_run = new_para.add_run(paragraph_text[0:index1])
                    paragraph_text = paragraph_text[index1:]

            #when the word is no longer found in the string we add the rest of the string as a separate run
            else:
                #new_doc.add_paragraph(paragraph_text)
                new_run = new_para.add_run(paragraph_text)
                #print(paragraph_text)
                paragraph_text = ""

    new_doc.save(filename + '.bold.docx')
    
bold_words('bacon.docx', 'hot')