#Chapter 17 of Automate the Boring Stuff by Al Sweigart
#https://automatetheboringstuff.com/3e/chapter17.html

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

def add_style_font_centering(paragraph, font_style, font_size, is_bold):
    run = paragraph.runs[0]
    run.font.name = font_style
    run.font.bold = is_bold
    run.font.size = Pt(font_size)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    """for run in opening_paragraph.runs:
        run.font.name = "Great Vibes"
        run.font.size = Pt(24)
    opening_paragraph_format = opening_paragraph.paragraph_format
    opening_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER"""

guest_file = 'guests.txt'
doc = docx.Document('template.docx')

#remove template paragraph (with the Great Vibes font)
p = doc.paragraphs[0]
p._element.getparent().remove(p._element)

with open(guest_file, encoding="UTF-8") as file_obj:
    for line in file_obj:
        opening_paragraph = doc.add_paragraph('It would be a pleasure to have the company of')
        add_style_font_centering(opening_paragraph, "Great Vibes", 24, False)

        name_paragraph = doc.add_paragraph(line.strip())
        add_style_font_centering(name_paragraph, "Calibri", 20, True)
        
        address_paragraph = doc.add_paragraph('at 11010 Memory Lane on the Evening of')
        add_style_font_centering(address_paragraph, "Great Vibes", 24, False)

        date_paragraph = doc.add_paragraph('April 1st')
        add_style_font_centering(date_paragraph, "Calibri", 20, False)

        time_paragraph = doc.add_paragraph("at 7 o'clock")
        add_style_font_centering(time_paragraph, "Great Vibes", 24, False)

        page_break = doc.add_paragraph()
        run = page_break.add_run()
        run.add_break(WD_BREAK.PAGE)
        
doc.save('invitation.docx')